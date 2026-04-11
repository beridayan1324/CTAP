import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def make_rtl(element):
    """Utility to force RTL direction on an element."""
    pPr = element._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def add_heading(doc, text, level):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = h.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    make_rtl(h)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    make_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    return p


def read_file_content(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"// Error reading {path}: {str(e)}"


def add_code_block(doc, code, filename):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"--- File: {filename} ---\n{code}")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 100)


def create_portfolio():
    print("Generating project portfolio...")
    doc = Document()

    # Define Project Details
    STUDENT_NAME = "[YOUR NAME]"
    TEACHER_NAME = "[TEACHER NAME]"
    SCHOOL_NAME = "[SCHOOL NAME]"
    GITHUB_LINK = "[YOUR GITHUB LINK]"
    SUBMISSION_DATE = "[DATE]"
    PROJECT_NAME = "CTAP - Communication Through Accessible Protocol"

    # --- 1. שער (Cover page) ---
    try:
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: pass

    add_paragraph(doc, "", size=24)
    p = add_paragraph(doc, "תיק פרויקט סייבר: " + PROJECT_NAME, bold=True, size=28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_paragraph(doc, "", size=24)
    add_paragraph(doc, f"מגיש: {STUDENT_NAME}", size=16)
    add_paragraph(doc, f"מנחה: {TEACHER_NAME}", size=16)
    add_paragraph(doc, f"בית ספר: {SCHOOL_NAME}", size=16)
    add_paragraph(doc, f"תאריך הגשה: {SUBMISSION_DATE}", size=16)
    add_paragraph(doc, f"קישור ל-GitHub: {GITHUB_LINK}", size=14)
    doc.add_page_break()

    # --- 2. תוכן עניינים ---
    add_heading(doc, "2. תוכן עניינים", 1)
    add_paragraph(doc, "הערה: יש לעדכן את תוכן העניינים בוורד לאחר הפקת המסמך (References -> Table of Contents)")
    
    sections = [
        "1. שער", "2. תוכן עניינים", "3. מבוא", "4. מבנה / ארכיטקטורה", 
        "5. מימוש הפרויקט", "6. מדריך למשתמש", "7. מדריך למפתח", 
        "8. סיכום אישי / רפלקציה", "9. ביבליוגרפיה", "10. נספחים"
    ]
    for s in sections:
        add_paragraph(doc, s, size=12)
    doc.add_page_break()


    # --- 3. מבוא ---
    add_heading(doc, "3. מבוא", 1)
    add_heading(doc, "רקע לפרויקט", 2)
    add_paragraph(doc, "במציאות של ימינו, לאנשים בעלי מוגבלויות פיזיות קשות (כגון שיתוק מלא או מחלות ניווניות כמו ALS) יש קושי עצום לתקשר עם העולם. פרויקט CTAP נועד לגשר על הפער הזה על ידי תרגום תנועות זעירות באצבעות הידיים לטקסט חיחבר. הפרויקט משלב חומרת אלקטרוניקה ולבישה, הצפנת נתונים חזקה ותקשורת רשת עכשווית כדי לאפשר ערוץ תקשורת פרטי ובטוח.", size=12)
    
    add_heading(doc, "מטרות המערכת", 2)
    add_paragraph(doc, "1. פיתוח אמצעי קלט חליפי דרך משקוף אצבעות וחיישני מגע/לחיצה.\n2. אבטחת ערוץ התקשורת באמצעות אלגוריתמים מתקדמים (AES-256-GCM).\n3. תיעוד אמין לשם אודיטינג (Audit Trail) דרך בסיס נתונים (SQLite).\n4. יצירת סביבת צ'אט מאובטחת בזמן אמת.", size=12)

    add_heading(doc, "תיאור המערכת", 2)
    add_paragraph(doc, "מערכת CTAP כוללת חומרה (ESP32 עם חיישנים), צד לקוח (Python/Node.js) לחיבור הסריאלי, ושרת מרכזי אשר מנהל את הצ'אט, מאמת זהויות באמצעות לחיצת-יד (Handshake) מבוססת סוד משותף, ומתעד נתונים מוצפנים למסד.", size=12)

    add_heading(doc, "אתגרים מרכזיים", 2)
    add_paragraph(doc, "1. סנכרון המידע בין החומרה (Arduino) סריאלית לשרת בזמן אמת.\n2. שילוב מנגנוני אבטחה קריפטוגרפיים (Nonce, Tags, AES-GCM) וניהול מפתחות בצד הלקוח והשרת.\n3. מניעת מתקפות הזרקה ורישום פעולות (Audit trail) שאינן פוגעות בפרטיות התוכן.", size=12)
    
    add_heading(doc, "חידושים טכנולוגיים", 2)
    add_paragraph(doc, "שימוש בהצפנה סימטרית מאומתת בזרם נתונים בזמן אמת, ממשק web מודרני מבוסס React המכיל 3D במקביל לניהול WebSockets ישירות מהדפדפן, ומעבר חלק לחלוטין בין תקשורת טורית מקומית לשרתי ענן מתקדמים.", size=12)

    doc.add_page_break()


    # --- 4. מבנה / ארכיטקטורה ---
    add_heading(doc, "4. מבנה / ארכיטקטורה", 1)
    
    add_heading(doc, "תיאור הטכנולוגיה", 2)
    tech_table = doc.add_table(rows=1, cols=4)
    tech_table.style = 'Table Grid'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = "רכיב/שכבה", "שפה/כלי", "ספריות עיקריות", "תפקיד"
    
    techs = [
        ("חומרה", "C++ / Arduino", "SerialPort", "עיבוד אותות חיישנים"),
        ("צד לקוח (מגשר)", "Python / Node.js", "websockets, cryptography", "הצפנה ותיווך רשת"),
        ("שרת (מרכזי)", "Node.js (Server) / Python", "ws, express, sqlite3", "ניהול Session ורישום"),
        ("צד לקוח (Web)", "React / Vite.js", "Three.js, Axios", "תצוגת ממשק המשתמש ממשק אודיט")
    ]
    for r in techs:
        row_cells = tech_table.add_row().cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = r

    add_heading(doc, "מסכי המערכת", 2)
    screen_table = doc.add_table(rows=1, cols=3)
    screen_table.style = 'Table Grid'
    hdr_cells = screen_table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = "מסך", "תיאור", "הרשאות"
    screens = [
        ("Landing", "מסך נחיתה תלת מימדי להמחשת הקידמה הטכנולוגית", "ציבורי"),
        ("Login / Register", "הזנת שם משתמש וסיסמה להתחברות והרשמה למערכת", "אורח"),
        ("Chat Room", "מרכז התקשורת של המערכת המשלב קלט חומרה והודעות", "משתמש מחובר"),
        ("Audit Logs", "הצגת יומני האבטחה (פעילות משתמשים ורשת)", "משתמש מחובר")
    ]
    for r in screens:
        row_cells = screen_table.add_row().cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text = r

    add_heading(doc, "אבטחה", 2)
    sec_table = doc.add_table(rows=1, cols=7)
    sec_table.style = 'Table Grid'
    hdr_cells = sec_table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text, hdr_cells[5].text, hdr_cells[6].text = "#", "מנגנון", "קובץ", "מחלקה/פונקציה", "תיאור והגנה", "איום שמונע", "חומרה"

    sec_data = [
        ("1", "הצפנה סימטרית מאומתת", "cryptoUtils.js", "encryptPayload()", "שימוש ב-AES-256-GCM כולל תווית אימות (Tag)", "האזנת סתר ושינוי נתונים", "קריטי"),
        ("2", "לחיצת יד מאובטחת", "server.js", "verifyHandshakeHash()", "אתגר ומענה (Challenge-Response) בשילוב סוד משותף", "התחזות (Spoofing) ולקוחות לא מורשים", "קריטי"),
        ("3", "גיבוב סיסמאות", "server.js", "bcrypt.hash()", "שמירת סיסמאות בגרסא מעורבלת עם מלח (Salt)", "גניבת מאגרי מידע (Data Breach)", "גבוה"),
        ("4", "נתיב ביקורת (Audit)", "database.js", "logMessage()", "רישום נתונים קריטיים, לרבות SHA256 של הודעות (ולא התווכן)", "הכחשה (Repudiation) ואובדן מידע", "גבוה"),
        ("5", "אימות JWT", "server.js", "verifyToken()", "אסימון גישה מתחדש ולא ניתן לזיוף", "חטיפת מושב (Session Hijacking)", "גבוה")
    ]
    for d in sec_data:
        row_cells = sec_table.add_row().cells
        for i in range(7): row_cells[i].text = d[i]

    doc.add_page_break()


    # --- 5. מימוש הפרויקט ---
    add_heading(doc, "5. מימוש הפרויקט", 1)
    
    # Analyze core files
    files_to_present = [
        {
            "name": "backend/server.js", 
            "desc": "ליבת השרת המטפלת בקריאות ה-REST, חיבורי יחידות דרך WebSocket ותהליכי לחיצת יד.",
            "path": r"c:\coding_stuff\CTAP\backend\server.js",
            "snippet_start": 110, "snippet_end": 165
        },
        {
            "name": "backend/cryptoUtils.js", 
            "desc": "פונקציות העזר הקריפטוגרפיות לטיפול ב-AES-GCM ואתגרי Handshake.",
            "path": r"c:\coding_stuff\CTAP\backend\cryptoUtils.js",
            "snippet_start": 8, "snippet_end": 44
        },
        {
            "name": "backend/database.js", 
            "desc": "ניהול מסד הנתונים SQLite לתיעוד משתמשים ויומני פעילות.",
            "path": r"c:\coding_stuff\CTAP\backend\database.js",
            "snippet_start": 15, "snippet_end": 48
        },
        {
            "name": "frontend/src/pages/Chat.jsx", 
            "desc": "מסך הצ'אט, כולל תקשורת עם ה-ESP32 מתוך הדפדפן באמצעות Web Serial API.",
            "path": r"c:\coding_stuff\CTAP\frontend\src\pages\Chat.jsx",
            "snippet_start": 125, "snippet_end": 148
        },
        {
            "name": "CTAP_FSR/CTAP_FSR.ino", 
            "desc": "קוד הבקר המורץ על גבי ה-ESP32 המתרגם לחיצות חיישנים לאותיות.",
            "path": r"c:\coding_stuff\CTAP\CTAP_FSR\CTAP_FSR.ino",
            "snippet_start": 40, "snippet_end": 64
        }
    ]

    for f in files_to_present:
        add_heading(doc, f'הקובץ: {f["name"]}', 2)
        add_paragraph(doc, f["desc"], size=12)
        content = read_file_content(f["path"])
        lines = content.split('\n')
        # Ensure slice indices are within range
        start = min(f["snippet_start"], len(lines))
        end = min(f["snippet_end"], len(lines))
        snippet = "\n".join(lines[start:end])
        add_code_block(doc, snippet, f["name"])
        add_paragraph(doc, "• הסבר על הקטע: הקוד לעיל ממחיש את המנגנון המרכזי שרלוונטי לקובץ. ניתן לראות כיצד מבוצעות הבדיקות השונות, הקמת האובייקטים והפנייה לפונקציות תכליתיות.", size=11, italic=True)

    doc.add_page_break()


    # --- 6. מדריך למשתמש ---
    add_heading(doc, "6. מדריך למשתמש", 1)
    
    add_heading(doc, "דרישות מערכת", 2)
    add_paragraph(doc, "- מחשב מריץ Windows / Linux.\n- התקנת Node.js 18+ או Python 3.10+.\n- דפדפן תומך מנוע כרומיום (Edge / Chrome) לשימוש ב-Web Serial API.\n- חומרת כפפת CTAP (בקר ESP32 + חיישנים) מחוברת ב-USB.", size=12)
    
    add_heading(doc, "תהליך עבודה (משתמש קצה)", 2)
    add_paragraph(doc, "1. הפעלת הדפדפן וכניסה למסך הבית (Landing).\n2. ביצוע הרשמה/התחברות בשילוב שם משתמש וסיסמה.\n3. לחיצה על כפתור Connect ESP32 לשם יצירת חיבור אל החומרה דרך ממשק ה-COM Port.\n4. תקתוק אצבעות על החיישן כדי לשלוח טקסט מבוסס בינארית למערכת.", size=12)

    doc.add_page_break()


    # --- 7. מדריך למפתח ---
    add_heading(doc, "7. מדריך למפתח", 1)
    add_paragraph(doc, "כדי לגשת לקוד המקור ולהרחיב את המערכת יש לבצע את השלבים הבאים:", size=12)
    code = '''git clone [repo_url]
cd CTAP

# Setup Backend
cd backend
npm install
node server.js

# Setup Frontend
cd frontend
npm install
npm run dev
'''
    add_code_block(doc, code, "Setup Script")
    add_heading(doc, "להוספת תכונות עתידיות (Features)", 2)
    add_paragraph(doc, "ניתן לשנות את מילון המושגים בתוך קובץ ה-C++. בקביעת אותיות חדשות מומלץ לבדוק את ההסתתרויות האפשריות. במסד הנתונים, ניתן להרחיב את מערכת ההרשאות על ידי מפתח נוסף בטבלת users.", size=12)
    
    doc.add_page_break()

    # --- 8. סיכום אישי / רפלקציה ---
    add_heading(doc, "8. סיכום אישי / רפלקציה", 1)
    add_paragraph(doc, "תהליך העבודה כלל אתגרים רבים החל משלב פיתוח החומרה, הרכבת המעגל החשמלי הפיזי והתמודדות עם רעשים באות (Noise). כדי לטפל בדברים אלו היה צורך לערוך מקצים של ניסוי וטעייה. תהליך הלמידה היה משמעותי: מעבודה עם WebSockets שמאפשרים עדכון ממשק ללא רענון עמוד, ועד ליישום מעשי של תקני הצפנה סימטרית AES256-GCM בסביבת פיתוח אמיתית.", size=12)
    
    add_paragraph(doc, "מסקנות בראייה לאחור: היה שווה להשקיע יותר זמן באפיון ובחירת פרוטוקולי התקשורת מתחילת הדרך, שכן החלפת Python ב-Node.js עבור יכולות Full Stack הייתה מאתגרת בהמשך התהליך.", size=12)
    
    add_paragraph(doc, "שיפורים נוכחיים ועתידיים: אינטגרציה מעמיקה יותר של מכונות גלגול מפתחות (Key Rotation), טיוב זמני התשה בחיבור הסריאלי ואף מעבר לחיבור אלחוטי ישיר (BLE / WiFi) במידה ויהיו משאבי סוללה עתידיים לציוד اللביש.", size=12)
    
    doc.add_page_break()

    # --- 9. ביבליוגרפיה ---
    add_heading(doc, "9. ביבליוגרפיה", 1)
    bl = [
        "1. תיעוד רשמי של Node.js - https://nodejs.org/docs",
        "2. תיעוד רשמי של React.js & Vite - https://reactjs.org | https://vitejs.dev",
        "3. הרחבת ממשק Web Serial - Mozilla Developer Network (MDN)",
        "4. מבוא לקריפטוגרפיה - NIST recommendations for Block Cipher Modes of Operation (GCM)",
        "5. תיעוד של ספריות Python Cryptography"
    ]
    for b in bl: add_paragraph(doc, b, size=12)
    doc.add_page_break()

    # --- 10. נספחים: הקוד המלא ---
    add_heading(doc, "10. נספחים", 1)
    add_paragraph(doc, "בחלק זה מוצג הקוד המלא המהווה את המערכת. כלל הקבצים, תוכן ומימוש מצורפים להלן לשם עמידה בהיקף ובדרישות ההגשה, ומונים את אלפי שורות הקוד שנכתבו בפרויקט כדי ליצור מערכת סייבר משולבת חומרה ותוכנה.", size=12)

    all_files = [
        r"c:\coding_stuff\CTAP\backend\server.js",
        r"c:\coding_stuff\CTAP\backend\database.js",
        r"c:\coding_stuff\CTAP\backend\cryptoUtils.js",
        r"c:\coding_stuff\CTAP\frontend\src\App.jsx",
        r"c:\coding_stuff\CTAP\frontend\src\pages\Chat.jsx",
        r"c:\coding_stuff\CTAP\frontend\src\pages\AuditLogs.jsx",
        r"c:\coding_stuff\CTAP\frontend\src\pages\Login.jsx",
        r"c:\coding_stuff\CTAP\frontend\src\pages\Register.jsx",
        r"c:\coding_stuff\CTAP\frontend\src\pages\Landing.jsx",
        r"c:\coding_stuff\CTAP\client\client.js",
        r"c:\coding_stuff\CTAP\CTAP_FSR\CTAP_FSR.ino",
        r"c:\coding_stuff\CTAP\CTAP_piezo\CTAP_piezo.ino",
        r"c:\coding_stuff\CTAP\client.py",
        r"c:\coding_stuff\CTAP\server.py",
        r"c:\coding_stuff\CTAP\index.html",
        r"c:\coding_stuff\CTAP\README.md",
        r"c:\coding_stuff\CTAP\readme.txt",
        r"c:\coding_stuff\CTAP\frontend\src\index.css",
        r"c:\coding_stuff\CTAP\frontend\src\App.css"
    ]

    for f_path in all_files:
        filename = os.path.basename(f_path)
        content = read_file_content(f_path)
        add_heading(doc, f"נספח: {filename}", 2)
        add_code_block(doc, content, filename)

    print("Saving document...")
    doc.save("תיק_פרויקט.docx")
    print("Done! Saved as תיק_פרויקט.docx")


if __name__ == "__main__":
    create_portfolio()
